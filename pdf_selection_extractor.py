#Basically, this shitcode script allows to run over the pdf files in a folder, 
#and to gather all the highlighted fragments into one summary word document called "summary" that will appear in the current directory
#:)

import docx
#pip install docx
#pip install lxml
#pip install Pillow
import fitz
#pip install fitz
#pip install pymupdf
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
#pip install os


fls = [] # will be a list of all files in the folder
pdfs = [] # will be a list of all .pdf files in the folder
for dirpath, dirnames, filenames in os.walk ('C:\\Users\\put\\your\\folder_with_pdfs_here'): #put your directory with pdfs here 
	fls.extend (filenames) #just all the files in the folder 
for fl in fls:
	if fl.endswith (('.pdf', '.PDF')):
		pdfs += [fl]

summary=docx.Document() #create the final word document titled "summary"

for p in pdfs: 
	po = fitz.open(p) #open pdf file
	mt = po.metadata #retrieve the metadata from pdf
	summary.add_paragraph ("Authors and title") #write into the summary for each pdf
	summary.add_paragraph (mt['author']) #author from metadata, if present
	summary.add_paragraph (mt['title']) #title from metadata, if present
	summary.add_paragraph ("Selections") #place for selections
	#summary.save("summary.docx") #save, though maybe not necessarily at this point

	#the one below finds the lighlighted fragments and adds 'em into the summary 
	pn = po.page_count
	for pgs in range (pn):
		page = po.load_page(pgs)
		for annot in page.annots():
			#print (annot)
			coord = annot.rect
			#print (coord)
			txtincoord = page.get_textbox(coord) 
			txtincoord2 = txtincoord.replace('\n', '') 
			summary.add_paragraph (str(txtincoord2)).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
			summary.save("summary.docx")
	summary.add_paragraph ('END OF THE ARTICLE\n')
	summary.save("summary.docx") #after writing the data from current pdf-file, the word file is saved and then goes iteration over next pdf  

